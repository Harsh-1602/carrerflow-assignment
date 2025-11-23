"""
Resume parsing tools for extracting text from PDF and DOCX files
"""
import PyPDF2
import pdfplumber
from docx import Document
from pathlib import Path
from typing import Dict, Optional
import json
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY


class ResumeParser:
    """Parse resume files and extract structured information"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF file using pdfplumber"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def parse_resume(file_path: str) -> Dict[str, any]:
        """
        Parse resume file and return structured information
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Dictionary with resume text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and parse
        extension = file_path.suffix.lower()
        
        if extension == ".pdf":
            text = ResumeParser.parse_pdf(str(file_path))
        elif extension == ".docx":
            text = ResumeParser.parse_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {extension}")
        
        # Extract basic structure (sections)
        sections = ResumeParser.extract_sections(text)
        
        return {
            "raw_text": text,
            "sections": sections,
            "file_name": file_path.name,
            "file_type": extension,
            "word_count": len(text.split())
        }
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """
        Extract common resume sections
        
        Args:
            text: Full resume text
            
        Returns:
            Dictionary of section names to content
        """
        sections = {}
        common_headers = [
            "SUMMARY", "OBJECTIVE", "PROFILE",
            "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT",
            "EDUCATION", "ACADEMIC BACKGROUND",
            "SKILLS", "TECHNICAL SKILLS", "COMPETENCIES",
            "PROJECTS", "PORTFOLIO",
            "CERTIFICATIONS", "CERTIFICATES",
            "AWARDS", "ACHIEVEMENTS", "HONORS"
        ]
        
        lines = text.split('\n')
        current_section = "HEADER"
        current_content = []
        
        for line in lines:
            line_upper = line.strip().upper()
            
            # Check if this line is a section header
            is_header = False
            for header in common_headers:
                if header in line_upper and len(line.strip()) < 50:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    
                    # Start new section
                    current_section = header
                    current_content = []
                    is_header = True
                    break
            
            if not is_header and line.strip():
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections


class DocumentGenerator:
    """Generate updated resume documents"""
    
    @staticmethod
    def create_text_file(content: str, output_path: str) -> str:
        """
        Create a text file with resume content
        
        Args:
            content: Resume text content
            output_path: Path for output file
            
        Returns:
            Path to created file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return str(output_path)
    
    @staticmethod
    def create_docx(content: str, output_path: str) -> str:
        """
        Create a DOCX file with resume content
        
        Args:
            content: Resume text content
            output_path: Path for output file
            
        Returns:
            Path to created file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # Add content paragraph by paragraph
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(str(output_path))
        return str(output_path)
    
    @staticmethod
    def create_pdf(content: str, output_path: str, title: str = "Resume") -> str:
        """
        Create a PDF file with resume content
        
        Args:
            content: Resume text content
            output_path: Path for output file
            title: Document title
            
        Returns:
            Path to created file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for PDF elements
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#1f77b4',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#2c3e50',
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=6
        )
        
        # Parse content and add to PDF
        lines = content.split('\n')
        section_headers = [
            "SUMMARY", "OBJECTIVE", "PROFILE",
            "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT",
            "EDUCATION", "ACADEMIC BACKGROUND",
            "SKILLS", "TECHNICAL SKILLS", "COMPETENCIES",
            "PROJECTS", "PORTFOLIO",
            "CERTIFICATIONS", "CERTIFICATES",
            "AWARDS", "ACHIEVEMENTS", "HONORS"
        ]
        
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1*inch))
                continue
            
            line_upper = line.upper()
            is_section = any(header in line_upper for header in section_headers)
            
            # Check if line is a section header
            if is_section and len(line) < 50:
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph(line, heading_style))
            else:
                # Regular content
                # Escape special characters for reportlab
                line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(line, body_style))
        
        # Build PDF
        doc.build(story)
        return str(output_path)
    
    @staticmethod
    def format_resume_sections(sections: Dict[str, str]) -> str:
        """
        Format resume sections into a cohesive document
        
        Args:
            sections: Dictionary of section names to content
            
        Returns:
            Formatted resume text
        """
        formatted = []
        
        # Define section order
        section_order = [
            "HEADER", "SUMMARY", "OBJECTIVE", "PROFILE",
            "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT",
            "EDUCATION", "ACADEMIC BACKGROUND",
            "SKILLS", "TECHNICAL SKILLS", "COMPETENCIES",
            "PROJECTS", "PORTFOLIO",
            "CERTIFICATIONS", "CERTIFICATES",
            "AWARDS", "ACHIEVEMENTS", "HONORS"
        ]
        
        # Add sections in order
        for section_name in section_order:
            if section_name in sections:
                formatted.append(f"\n{section_name}\n{'=' * len(section_name)}")
                formatted.append(sections[section_name])
        
        # Add any remaining sections not in the predefined order
        for section_name, content in sections.items():
            if section_name not in section_order:
                formatted.append(f"\n{section_name}\n{'=' * len(section_name)}")
                formatted.append(content)
        
        return '\n'.join(formatted)
